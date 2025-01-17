'''
Resume parser to give information to be evaluated in reviewer.py
'''
from docx import Document
from nltk.sentiment.vader import SentimentIntensityAnalyzer

PAGE_COUNT_ALLOWED = True
try:
    import win32com.client as win32
except ModuleNotFoundError:
    PAGE_COUNT_ALLOWED = False
    print('Page Count will not work in this system. Most likely cause of error is that the current Operating System is macOS (the library for the page count functions are not supported on macOS). If you need to use the page count functionality, consider moving to a Windows system or using a Virtual Machine or dual-booting using Bootcamp.')

from resumeAnalyzer import spellCheck
import re


class Resume:
    '''
    Loads a resume Word document and parses it to automate some basic resume reviewing tasks
    '''

    def __init__(self, file_path):
        self.PATH = file_path
        
        self.document = Document(self.PATH) # load word document

        self.content = []
        # add each paragraph text to content list
        for paragraph in self.document.paragraphs:
            self.content.append(paragraph.text)
    
    def page_count(self):
        '''
        Returns the number of pages in the document/resume
        '''
        try:
            # open word document
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(self.PATH)
            
            # get number of sheets
            doc.Repaginate()
            pages = doc.ComputeStatistics(2)

            # close word
            doc.Close()
            word.Quit()

            return pages
        
        # if error while loading document or getting page count
        except:
            print('[ERROR] Error Checking Page Count. Make sure document is not in Protected View Mode.')
        
        return None
    
    def get_text(self):
        '''
        Parses the content list to refine the analyzable text
        '''
        # Remove short words and phrases that aren't narrative
        for c in reversed(self.content):
            if len(c) < 5 or len(c.strip().split(' ')) < 5:
                self.content.remove(c)
        
        # Delete escape characters
        self.text = []
        for c in self.content:
            self.text.append(c.replace('\t', ''))
        
        return self.text
    
    def text_sentiment(self):
        '''
        Returns the calculated max sentiment of the narrative text
        '''
        nltk_sentiment = SentimentIntensityAnalyzer()

        # Make sure self.text actually has text in it
        try:
            self.text += ''
        except AttributeError:
            self.get_text()

        # Get sentiment values for each sentence
        neg, neu, pos, count = (0, 0, 0, 0)
        for paragraph in self.text:
            for sentence in paragraph.split('.'):
                sentiment_results = nltk_sentiment.polarity_scores(sentence.strip()) # get sentiment values for sentence
                count += 1
                neg += sentiment_results['neg']
                neu += sentiment_results['neu']
                pos += sentiment_results['pos']
        
        # Find maximum sentiment value
        sentiment_list = {'negative': neg/count, 'neutral': neu/count, 'positive': pos/count}
        maximum = max([neg/count, neu/count, pos/count])

        for _type, value in sentiment_list.items():
            if value == maximum:
                sentiment = (_type, value)

        return sentiment

    def spell_check(self):
        '''
        Corrects misspelled words in the Word document
        '''
        mistakes = 0

        # iterate over every word in the document; replace it if it is mispelled
        for paragraph in self.document.paragraphs:
            for word in spellCheck.words(paragraph.text):
                corrected_word = spellCheck.correction(word)
                if not bool(re.search(r'\d', word)) and corrected_word != word:
                    paragraph.text = paragraph.text.replace(word, corrected_word)
                    mistakes += 1
                    
        self.document.save(self.PATH)
        return mistakes


if __name__ == "__main__":
    import os

    if PAGE_COUNT_ALLOWED == True:
        resume = Resume(os.path.dirname(os.path.realpath(__file__)) + '\\test_data\\resume_template.docx')
    else:
        resume = Resume(os.path.dirname(os.path.realpath(__file__)) + '/test_data/resume_template.docx')
    resume.get_text()
    if PAGE_COUNT_ALLOWED == True:
        print(f'Detected {resume.page_count()} pages in resume')
    sentiment = resume.text_sentiment()
    print(f'Detected tone of resume is {sentiment[0]} of value {sentiment[1]}')
    print(f'Corrected {resume.spell_check()} misspelled words in the resume')